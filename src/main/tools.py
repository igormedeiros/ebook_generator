from crewai_tools import Tool
from docx import Document

class ReplaceTextInDocxTool(Tool):
    def __init__(self):
        super().__init__()

    def run(self, input_file, output_file, replacements):
        """
        Substitui textos em um arquivo .docx e salva o resultado em um novo arquivo.
        
        Args:
        - input_file (str): Caminho para o arquivo .docx de entrada.
        - output_file (str): Caminho para o arquivo .docx de saída.
        - replacements (dict): Dicionário onde a chave é o texto antigo e o valor é o texto novo.
        
        Returns:
        - str: Mensagem de sucesso ou erro.
        """
        try:
            # Abrir o documento .docx existente
            doc = Document(input_file)

            # Iterar sobre os parágrafos e substituir o texto
            for para in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in para.text:
                        para.text = para.text.replace(old_text, new_text)

            # Iterar sobre as tabelas, se houver
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for old_text, new_text in replacements.items():
                            if old_text in cell.text:
                                cell.text = cell.text.replace(old_text, new_text)

            # Salvar o documento atualizado
            doc.save(output_file)

            return f"Substituição de texto completa. Arquivo salvo como {output_file}."
        except Exception as e:
            return f"Erro ao processar o arquivo: {str(e)}"
